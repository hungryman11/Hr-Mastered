from io import BytesIO
from pathlib import Path
from datetime import datetime

from django.conf import settings
from django.core.exceptions import ValidationError
from django.db import transaction
from django.utils import timezone

from core.models import ApprovalDocument, ApprovalDecision, Employee, EmployeeRole, LeaveApprovalStep, LeaveBalance, LeaveRequest, LeaveType, OrgUnit, OnboardingStatus, LeaveApprovalPolicy, ApprovalDelegation
from zoho.services import ZohoWorkDriveService


class ApprovalRoutingService:
    @staticmethod
    def get_leave_approvers(employee: Employee):
        return [step.approver for step in ApprovalRoutingService.build_route(employee)]

    @staticmethod
    def build_route(employee: Employee):
        """Return the configured Department Head -> HR route.

        People sharing a stage approve in parallel; every person in that stage
        must approve before the next stage is notified.
        """
        if not employee.company_id:
            raise ValidationError('Employee must belong to a company.')

        active = Employee.objects.filter(company=employee.company, is_active=True, deleted_at__isnull=True)
        route = []

        def find_policy(unit: OrgUnit):
            visited = set()
            while unit and unit.pk not in visited:
                visited.add(unit.pk)
                try:
                    return unit.leave_policy
                except LeaveApprovalPolicy.DoesNotExist:
                    unit = unit.parent
            return None

        def apply_delegation(approver: Employee):
            # Find an active delegation where approver is delegated to someone else today in the same company.
            from django.utils import timezone as _tz
            today = _tz.now().date()
            deleg = ApprovalDelegation.objects.filter(
                approver=approver,
                active=True,
                start_date__lte=today,
                end_date__gte=today,
                delegate_to__company=approver.company,
                delegate_to__is_active=True,
                delegate_to__deleted_at__isnull=True,
            ).first()
            return deleg.delegate_to if deleg else approver

        dept_heads = []
        org_unit = employee.org_unit
        policy = find_policy(org_unit) if org_unit else None
        if policy:
            if policy.first_approver_type == LeaveApprovalPolicy.ApproverType.MANAGER:
                if employee.manager and employee.manager.is_active and employee.manager_id != employee.pk:
                    dept_heads = [apply_delegation(employee.manager)]
            elif policy.first_approver_type == LeaveApprovalPolicy.ApproverType.SPECIFIC:
                emp = policy.first_approver_employee
                if emp and emp.is_active and emp.pk != employee.pk:
                    dept_heads = [apply_delegation(emp)]
            else:
                unit = org_unit
                visited_units = set()
                while unit and unit.pk not in visited_units:
                    visited_units.add(unit.pk)
                    if unit.head_id and unit.head_id != employee.pk and unit.head.is_active:
                        dept_heads.append(apply_delegation(unit.head))
                    unit = unit.parent
        else:
            # Walk the org-unit ancestor chain and stop at the first active head found.
            # Only one department-head approver is needed for Stage 1; collecting every
            # ancestor head would make the route multi-level rather than 2-stage.
            unit = org_unit
            visited_units = set()
            while unit and unit.pk not in visited_units:
                visited_units.add(unit.pk)
                if unit.head_id and unit.head_id != employee.pk and unit.head.is_active:
                    dept_heads = [apply_delegation(unit.head)]
                    break  # stop at the closest head – this is Stage 1
                unit = unit.parent
            if not dept_heads and employee.manager and employee.manager.is_active and employee.manager_id != employee.pk:
                dept_heads = [apply_delegation(employee.manager)]

        if not dept_heads:
            raise ValidationError("Leave approval cannot proceed because no department head is configured for this employee's department.")

        seen = set()
        unique_dept_heads = []
        for h in dept_heads:
            if h.pk not in seen and h.pk != employee.pk:
                seen.add(h.pk)
                unique_dept_heads.append(h)
        if not unique_dept_heads:
            raise ValidationError("Leave approval cannot proceed because no department head is configured for this employee's department.")

        route.append((LeaveApprovalStep.Stage.DEPT_HEAD, unique_dept_heads))

        hr_admins = list(active.filter(role=EmployeeRole.HR_ADMIN).exclude(pk=employee.pk).order_by('id'))
        if not hr_admins:
            raise ValidationError('No active HR administrator is configured for this company.')
        hr_admins = [apply_delegation(h) for h in hr_admins]
        seen = set()
        unique_hr_admins = []
        for h in hr_admins:
            if h.pk not in seen and h.pk != employee.pk:
                seen.add(h.pk)
                unique_hr_admins.append(h)
        if not unique_hr_admins:
            raise ValidationError('No active HR administrator is configured for this company.')

        route.append((LeaveApprovalStep.Stage.HR, unique_hr_admins))

        return [
            LeaveApprovalStep(sequence=sequence, stage=stage, approver=approver, company=employee.company)
            for sequence, (stage, approvers) in enumerate(route, start=1)
            for approver in {person.pk: person for person in approvers}.values()
        ]

    @staticmethod
    def create_steps(leave_request: LeaveRequest, approval_round=None):
        steps = ApprovalRoutingService.build_route(leave_request.employee)
        for step in steps:
            step.leave_request = leave_request
            step.approval_round = approval_round or leave_request.approval_round
            step.created_by = leave_request.employee
            step.updated_by = leave_request.employee
        LeaveApprovalStep.objects.bulk_create(steps)
        return steps

    @staticmethod
    def get_current_steps(leave_request: LeaveRequest):
        pending = leave_request.approval_steps.filter(
            approval_round=leave_request.approval_round,
            status=LeaveApprovalStep.Status.PENDING,
        )
        first = pending.order_by('sequence').first()
        if not first:
            return leave_request.approval_steps.none()
        return pending.filter(sequence=first.sequence).select_related('approver')

    @staticmethod
    def next_approver(employee: Employee):
        route = ApprovalRoutingService.build_route(employee)
        return route[0].approver if route else None


class ApprovalDocumentService:
    @staticmethod
    def _document_dir():
        base_dir = Path(getattr(settings, 'APPROVAL_DOCUMENT_DIR', Path(settings.BASE_DIR) / 'generated_documents'))
        base_dir.mkdir(parents=True, exist_ok=True)
        return base_dir

    @staticmethod
    def _write_docx(file_path: Path, title: str, lines: list[str]):
        try:
            from docx import Document
        except ImportError as exc:
            raise ValidationError('python-docx is required to generate approval documents.') from exc

        document = Document()
        document.add_heading(title, 0)
        for line in lines:
            document.add_paragraph(line)
        document.save(str(file_path))

    @staticmethod
    def create_for_leave_request(leave_request: LeaveRequest, document_type: str, actor: Employee):
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        file_name = f"{leave_request.employee.username}_{leave_request.leave_type.name}_{document_type.lower()}_{timestamp}.docx".replace(' ', '_')
        file_path = ApprovalDocumentService._document_dir() / file_name
        title = f"Leave {document_type.title()}"
        lines = [
            f"Company: {leave_request.company.name}",
            f"Employee: {leave_request.employee.get_full_name() or leave_request.employee.username}",
            f"Leave Type: {leave_request.leave_type.name}",
            f"Start Date: {leave_request.start_date}",
            f"End Date: {leave_request.end_date}",
            f"Days Requested: {leave_request.days_requested}",
            f"Status: {leave_request.status}",
            f"Reviewed By: {actor.get_full_name() or actor.username}",
            f"Reviewed At: {timezone.now()}",
            f"Reason: {leave_request.reason or '-'}",
            f"Rejection Reason: {leave_request.rejection_reason or '-'}",
        ]
        ApprovalDocumentService._write_docx(file_path, title, lines)
        approval_doc = ApprovalDocument.objects.create(
            company=leave_request.company,
            leave_request=leave_request,
            document_type=document_type,
            file_name=file_name,
            file_path=str(file_path),
            created_by=actor,
        )
        from core.delivery import DeliveryService
        DeliveryService.enqueue_approval_document(approval_doc)

        return approval_doc


class LeaveRequestDocumentService:
    """Creates the standard WorkDrive leave form from the employee's answers."""

    @staticmethod
    def _count_working_days(start_date, end_date) -> int:
        """Count Monday–Friday working days between two dates (inclusive)."""
        from datetime import timedelta
        total = 0
        current = start_date
        while current <= end_date:
            if current.weekday() < 5:  # Mon=0 … Fri=4
                total += 1
            current += timedelta(days=1)
        return total

    @staticmethod
    def generate(leave_request: LeaveRequest) -> tuple[str, bytes]:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise ValidationError('python-docx is required to generate leave documents.') from exc

        working_days = LeaveRequestDocumentService._count_working_days(
            leave_request.start_date, leave_request.end_date
        )

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        heading = document.add_heading('Leave Request Form', 0)
        heading.runs[0].font.size = Pt(18)
        document.add_paragraph(f'Company: {leave_request.company.name}')
        document.add_paragraph(f'Request reference: {leave_request.uuid}')

        table = document.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        fields = (
            ('Employee', leave_request.employee.get_full_name() or leave_request.employee.username),
            ('Leave type', leave_request.leave_type.name),
            ('Start date', str(leave_request.start_date)),
            ('End date', str(leave_request.end_date)),
            ('Calendar days', str((leave_request.end_date - leave_request.start_date).days + 1)),
            ('Working days', str(working_days)),
            ('Days requested', f'{leave_request.days_requested} day(s)'),
            ('Reason for leave', leave_request.reason or '-'),
            ('Contact during leave', leave_request.contact_during_leave),
            ('Emergency contact', f'{leave_request.emergency_contact_name} ({leave_request.emergency_contact_phone})'),
            ('Handover contact', leave_request.handover_contact),
            ('Handover notes', leave_request.handover_notes),
        )
        for label, value in fields:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)

        # Approval history — initially empty; re-generated after each decision
        document.add_heading('Approval History', level=1)
        decisions = list(
            leave_request.approval_decisions.select_related('actor')
            .order_by('decided_at')
        )
        if decisions:
            hist_table = document.add_table(rows=1, cols=4)
            hist_table.style = 'Table Grid'
            hdr = hist_table.rows[0].cells
            hdr[0].text = 'Stage'
            hdr[1].text = 'Decided by'
            hdr[2].text = 'Decision'
            hdr[3].text = 'Date'
            for dec in decisions:
                row_cells = hist_table.add_row().cells
                row_cells[0].text = dec.stage or '-'
                row_cells[1].text = dec.actor.get_full_name() or dec.actor.username
                row_cells[2].text = dec.decision
                row_cells[3].text = str(dec.decided_at.date()) if dec.decided_at else '-'
        else:
            document.add_paragraph('No decisions recorded yet.')

        document.add_paragraph(
            'This form was generated automatically by the HR platform. '
            f'Current status: {leave_request.status}.'
        )

        output = BytesIO()
        document.save(output)
        file_name = f'leave_request_{leave_request.uuid}.docx'
        return file_name, output.getvalue()


class OnboardingService:
    @staticmethod
    def onboard_employee(employee: Employee, created_by: Employee = None):
        if not employee.company_id:
            raise ValueError('Employee must belong to a company before onboarding.')

        if getattr(employee, 'workdrive_folder', None):
            if employee.onboarding_status != OnboardingStatus.COMPLETE:
                employee.onboarding_status = OnboardingStatus.COMPLETE
                employee.save(update_fields=['onboarding_status', 'updated_at'])
            return employee.workdrive_folder

        service = ZohoWorkDriveService()
        folder = service.create_folder(
            company=employee.company,
            folder_name=f'{employee.get_full_name() or employee.username} Files',
            created_by=created_by,
            employee=employee,
        )
        employee.onboarding_status = OnboardingStatus.COMPLETE
        employee.save(update_fields=['onboarding_status', 'updated_at'])
        return folder


class LeaveService:
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

    @staticmethod
    def calculate_working_days(start_date, end_date, company=None):
        from datetime import timedelta
        from core.holidays import nigerian_public_holidays
        from core.models import CompanyHoliday, CompanyWorkCalendar
        working_weekdays = {0, 1, 2, 3, 4}
        include_nigerian_public_holidays = True
        if company:
            calendar = CompanyWorkCalendar.objects.filter(company=company).first()
            if calendar:
                working_weekdays = set(calendar.working_weekdays)
                include_nigerian_public_holidays = calendar.include_nigerian_public_holidays
        holidays = nigerian_public_holidays(start_date.year, end_date.year) if include_nigerian_public_holidays else set()
        if company:
            holidays.update(CompanyHoliday.objects.filter(
                company=company, date__range=(start_date, end_date),
            ).values_list('date', flat=True))
        total = 0
        current = start_date
        while current <= end_date:
            if current.weekday() in working_weekdays and current not in holidays:
                total += 1
            current += timedelta(days=1)
        return total

    @staticmethod
    def validate_document(document_file):
        """Validate filename, declared MIME type, size, and lightweight file signature.

        Returns None immediately when no file is supplied (supporting documents are optional).
        """
        if not document_file:
            return None
        original_name = getattr(document_file, 'name', '')
        safe_name = Path(original_name).name
        if not safe_name or safe_name != original_name or safe_name in {'.', '..'}:
            raise ValidationError('Document filename is invalid.')
        ext = Path(safe_name).suffix.lower()
        if ext not in LeaveService.ALLOWED_EXTENSIONS:
            raise ValidationError(f"File extension '{ext}' is not supported. Allowed: {', '.join(sorted(LeaveService.ALLOWED_EXTENSIONS))}")
        size = getattr(document_file, 'size', 0)
        if size <= 0:
            raise ValidationError('Attached document cannot be empty.')
        if size > LeaveService.MAX_FILE_SIZE:
            raise ValidationError('Attached document exceeds the 10 MB size limit.')

        expected_types = {
            '.pdf': {'application/pdf'},
            '.doc': {'application/msword'},
            '.docx': {'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
            '.png': {'image/png'},
            '.jpg': {'image/jpeg'}, '.jpeg': {'image/jpeg'},
        }
        content_type = getattr(document_file, 'content_type', '')
        if content_type and content_type not in expected_types[ext]:
            raise ValidationError('Document content type does not match its extension.')
        header = document_file.read(8)
        document_file.seek(0)
        signatures = {
            '.pdf': b'%PDF-', '.docx': b'PK\\x03\\x04', '.png': b'\\x89PNG\\r\\n\\x1a\\n',
            '.jpg': b'\\xff\\xd8\\xff', '.jpeg': b'\\xff\\xd8\\xff',
        }
        signature = signatures.get(ext)
        if signature and not header.startswith(signature):
            raise ValidationError('Document content does not match its extension.')
        return safe_name

    @staticmethod
    def _notify_steps(leave_request: LeaveRequest, steps):
        from core.delivery import DeliveryService
        for step in steps:
            if not step.approver.email:
                continue
            DeliveryService.enqueue_email(
                company=leave_request.company,
                recipient=step.approver.email,
                subject=f"[Leave Request] Action required: {step.stage}",
                body=(f"A leave request from {leave_request.employee.get_full_name() or leave_request.employee.username} "
                      f"is awaiting your {step.stage} approval.\n\n"
                      f"Duration: {leave_request.start_date} to {leave_request.end_date} "
                      f"({leave_request.days_requested} days)\n"
                      f"Reason: {leave_request.reason or 'N/A'}"),
                template_name='leave_approval_action_required',
                sent_by=leave_request.employee,
            )

    @staticmethod
    def request_leave(
        *,
        employee: Employee,
        leave_type: LeaveType,
        start_date,
        end_date,
        days_requested,
        reason='',
        contact_during_leave='',
        emergency_contact_name='',
        emergency_contact_phone='',
        handover_contact='',
        handover_notes='',
        document_file=None
    ):
        if employee.company_id != leave_type.company_id:
            raise ValidationError('Leave type must belong to the same company as the employee.')

        if start_date and end_date and end_date < start_date:
            raise ValidationError({'end_date': 'End date must be on or after start date.'})

        calculated_days = LeaveService.calculate_working_days(start_date, end_date, employee.company)
        if calculated_days <= 0:
            raise ValidationError({'start_date': 'Leave must include at least one working day.'})
        max_days = leave_type.max_days_per_request if leave_type.max_days_per_request is not None else settings.MAX_LEAVE_WORKING_DAYS
        if calculated_days > max_days:
            raise ValidationError({
                'end_date': f'Leave requests for {leave_type.name} may not exceed {max_days} working days.'
            })
        days_requested = calculated_days

        supporting_name = None
        if document_file:
            supporting_name = LeaveService.validate_document(document_file)
            if supporting_name:
                document_file.seek(0)
        elif leave_type.requires_supporting_document:
            raise ValidationError(
                f"A supporting document is required for '{leave_type.name}' leave requests."
            )

        # Check for overlapping leave requests
        overlap_exists = LeaveRequest.objects.filter(
            employee=employee,
            status__in=[
                LeaveRequest.Status.PENDING,
                LeaveRequest.Status.PENDING_DEPARTMENT_HEAD,
                LeaveRequest.Status.PENDING_HR,
                LeaveRequest.Status.APPROVED,
            ],
            start_date__lte=end_date,
            end_date__gte=start_date,
        ).exists()
        if overlap_exists:
            raise ValidationError('You already have a pending or approved leave request during these dates.')

        with transaction.atomic():
            leave_request = LeaveRequest.objects.create(
                company=employee.company,
                employee=employee,
                leave_type=leave_type,
                start_date=start_date,
                end_date=end_date,
                days_requested=days_requested,
                reason=reason,
                contact_during_leave=contact_during_leave,
                emergency_contact_name=emergency_contact_name,
                emergency_contact_phone=emergency_contact_phone,
                handover_contact=handover_contact,
                handover_notes=handover_notes,
                status=LeaveRequest.Status.PENDING_DEPARTMENT_HEAD,
                created_by=employee,
                updated_by=employee,
            )

            document_upload_error = None
            generated_name = None
            try:
                folder = OnboardingService.onboard_employee(employee, created_by=employee)
                service = ZohoWorkDriveService()
                generated_name, generated_content = LeaveRequestDocumentService.generate(leave_request)
                generated_document = service.upload_document(
                    employee=employee, folder=folder, document_name=generated_name,
                    document_type='docx', file_content=generated_content, uploaded_by=employee,
                )
                leave_request.document_name = generated_name
                leave_request.zoho_file_id = generated_document.zoho_file_id
                leave_request.workdrive_url = f"https://workdrive.zoho.com/file/{generated_document.zoho_file_id}"

                if document_file:
                    supporting_document = service.upload_document(
                        employee=employee,
                        folder=folder,
                        document_name=supporting_name,
                        document_type=Path(supporting_name).suffix.lstrip('.'),
                        file_content=document_file.read(),
                        uploaded_by=employee,
                    )
                    leave_request.supporting_document_name = supporting_name
                    leave_request.supporting_zoho_file_id = supporting_document.zoho_file_id
                    leave_request.supporting_workdrive_url = f"https://workdrive.zoho.com/file/{supporting_document.zoho_file_id}"
                leave_request.save(update_fields=[
                    'document_name', 'zoho_file_id', 'workdrive_url',
                    'supporting_document_name', 'supporting_zoho_file_id', 'supporting_workdrive_url',
                    'updated_at',
                ])
            except ValidationError:
                raise
            except Exception as exc:
                document_upload_error = exc
                if generated_name:
                    leave_request.document_name = generated_name
                else:
                    leave_request.document_name = f"leave_request_{leave_request.uuid}.docx"
                leave_request.zoho_file_id = None
                leave_request.workdrive_url = None
                leave_request.supporting_document_name = supporting_name
                leave_request.supporting_zoho_file_id = None
                leave_request.supporting_workdrive_url = None
                leave_request.save(update_fields=[
                    'document_name', 'zoho_file_id', 'workdrive_url',
                    'supporting_document_name', 'supporting_zoho_file_id', 'supporting_workdrive_url',
                    'updated_at',
                ])

            ApprovalRoutingService.create_steps(leave_request)
            LeaveService._notify_steps(leave_request, ApprovalRoutingService.get_current_steps(leave_request))

            if document_upload_error:
                leave_request.rejection_reason = (
                    'WorkDrive document upload failed; leave request was created but document metadata is pending.'
                )
                leave_request.save(update_fields=['rejection_reason', 'updated_at'])

            return leave_request

    @staticmethod
    def approve_leave(leave_request: LeaveRequest, reviewed_by: Employee, reason: str):
        if reviewed_by.company_id != leave_request.company_id:
            raise ValidationError('Reviewer must belong to the same company.')

        from core.delivery import DeliveryService

        if leave_request.status not in (
            LeaveRequest.Status.PENDING_DEPARTMENT_HEAD,
            LeaveRequest.Status.PENDING_HR,
            LeaveRequest.Status.PENDING,
        ):
            raise ValidationError('Only a pending leave request can be approved.')
        if not reason or not reason.strip():
            raise ValidationError({'reason': 'An approval reason is required.'})

        with transaction.atomic():
            # Lock the leave_request and its pending approval steps to avoid races
            leave_request = LeaveRequest.objects.select_for_update().get(pk=leave_request.pk)
            current_steps = ApprovalRoutingService.get_current_steps(leave_request).select_for_update()
            step = current_steps.filter(approver=reviewed_by).first()
            if not step:
                raise ValidationError('You are not assigned to the current approval stage.')
            current_sequence = step.sequence
            if step.status != LeaveApprovalStep.Status.PENDING:
                # Defensive check: guard against replayed decisions
                raise ValidationError('This approval step has already been decided.')

            step.status = LeaveApprovalStep.Status.APPROVED
            step.decision_reason = reason.strip()
            step.decided_at = timezone.now()
            step.updated_by = reviewed_by
            step.save(update_fields=['status', 'decision_reason', 'decided_at', 'updated_by', 'updated_at'])

            # Record immutable approval decision
            # Prevent duplicate decisions by same actor for same step
            if ApprovalDecision.objects.filter(leave_request=leave_request, approval_step=step, actor=reviewed_by).exists():
                raise ValidationError('This approver has already made a decision for this step.')
            ApprovalDecision.objects.create(
                company=leave_request.company,
                leave_request=leave_request,
                approval_step=step,
                actor=reviewed_by,
                stage=step.stage,
                sequence=step.sequence,
                decision=ApprovalDecision.Decision.APPROVED,
                reason=reason.strip(),
                decided_at=step.decided_at,
                created_by=reviewed_by,
                updated_by=reviewed_by,
            )

            # If there are still pending approvers at this sequence, do not advance
            if leave_request.approval_steps.filter(sequence=current_sequence, status=LeaveApprovalStep.Status.PENDING).exists():
                return leave_request

            # Determine next steps (if any) after current sequence
            next_steps = ApprovalRoutingService.get_current_steps(leave_request)
            if next_steps.exists():
                next_stage = next_steps.first().stage
                if next_stage == LeaveApprovalStep.Stage.HR:
                    leave_request.status = LeaveRequest.Status.PENDING_HR
                    leave_request.save(update_fields=['status', 'updated_at'])
                LeaveService._notify_steps(leave_request, next_steps)
                # Re-generate and re-upload the leave form to keep WorkDrive's copy current
                try:
                    folder = OnboardingService.onboard_employee(leave_request.employee, created_by=reviewed_by)
                    gen_name, gen_content = LeaveRequestDocumentService.generate(leave_request)
                    ZohoWorkDriveService().upload_document(
                        employee=leave_request.employee, folder=folder,
                        document_name=gen_name, document_type='docx',
                        file_content=gen_content, uploaded_by=reviewed_by,
                    )
                except Exception:
                    pass  # Non-blocking: WorkDrive sync failures must not roll back the approval
                return leave_request

            # Final approval stage: check and deduct balance under DB lock
            balance = LeaveBalance.objects.select_for_update().filter(
                company=leave_request.company,
                employee=leave_request.employee,
                leave_type=leave_request.leave_type
            ).first()
            if not balance:
                # Create a balance row if none exists. Creating inside the transaction avoids race, then re-lock.
                balance = LeaveBalance.objects.create(
                    company=leave_request.company,
                    employee=leave_request.employee,
                    leave_type=leave_request.leave_type,
                    allocated_days=leave_request.leave_type.default_days,
                    used_days=0,
                    created_by=reviewed_by,
                    updated_by=reviewed_by,
                )
                balance = LeaveBalance.objects.select_for_update().get(pk=balance.pk)

            if balance.remaining_days < leave_request.days_requested:
                raise ValidationError('Insufficient leave balance to approve this request.')

            # Commit approval to leave request and deduct balance
            leave_request.status = LeaveRequest.Status.APPROVED
            leave_request.reviewed_by = reviewed_by
            leave_request.reviewed_at = timezone.now()
            leave_request.updated_by = reviewed_by
            leave_request.save(update_fields=['status', 'reviewed_by', 'reviewed_at', 'updated_by', 'updated_at'])

            balance.used_days += leave_request.days_requested
            balance.updated_by = reviewed_by
            balance.save(update_fields=['used_days', 'updated_by', 'updated_at'])

            ApprovalDocumentService.create_for_leave_request(leave_request, ApprovalDocument.DocumentType.APPROVAL, reviewed_by)

            # Notify employee of approval
            if leave_request.employee.email:
                subject = f"[Leave Approved] Your {leave_request.leave_type.name} request has been approved"
                body = (
                    f"Hello {leave_request.employee.get_full_name() or leave_request.employee.username},\n\n"
                    f"Your leave request for {leave_request.leave_type.name} from {leave_request.start_date} to {leave_request.end_date} has been APPROVED by {reviewed_by.get_full_name() or reviewed_by.username}.\n"
                )
                DeliveryService.enqueue_email(
                    company=leave_request.company,
                    recipient=leave_request.employee.email,
                    subject=subject,
                    body=body,
                    template_name='leave_approved_notification',
                    sent_by=reviewed_by,
                )

            return leave_request

    @staticmethod
    def reject_leave(leave_request: LeaveRequest, reviewed_by: Employee, reason=''):
        if reviewed_by.company_id != leave_request.company_id:
            raise ValidationError('Reviewer must belong to the same company.')

        from core.delivery import DeliveryService

        if leave_request.status not in (
            LeaveRequest.Status.PENDING_DEPARTMENT_HEAD,
            LeaveRequest.Status.PENDING_HR,
            LeaveRequest.Status.PENDING,
        ):
            raise ValidationError('Only a pending leave request can be rejected.')
        if not reason or not reason.strip():
            raise ValidationError({'reason': 'A rejection reason is required.'})

        with transaction.atomic():
            leave_request = LeaveRequest.objects.select_for_update().get(pk=leave_request.pk)
            step = ApprovalRoutingService.get_current_steps(leave_request).select_for_update().filter(approver=reviewed_by).first()
            if not step:
                raise ValidationError('You are not assigned to the current approval stage.')
            step.status = LeaveApprovalStep.Status.REJECTED
            step.decision_reason = reason.strip()
            step.decided_at = timezone.now()
            step.updated_by = reviewed_by
            step.save(update_fields=['status', 'decision_reason', 'decided_at', 'updated_by', 'updated_at'])

            # Record immutable rejection decision
            if ApprovalDecision.objects.filter(leave_request=leave_request, approval_step=step, actor=reviewed_by).exists():
                raise ValidationError('This approver has already made a decision for this step.')
            ApprovalDecision.objects.create(
                company=leave_request.company,
                leave_request=leave_request,
                approval_step=step,
                actor=reviewed_by,
                stage=step.stage,
                sequence=step.sequence,
                decision=ApprovalDecision.Decision.REJECTED,
                reason=reason.strip(),
                decided_at=step.decided_at,
                created_by=reviewed_by,
                updated_by=reviewed_by,
            )

            leave_request.status = LeaveRequest.Status.REJECTED
            leave_request.reviewed_by = reviewed_by
            leave_request.reviewed_at = timezone.now()
            leave_request.rejection_reason = reason.strip()
            leave_request.updated_by = reviewed_by
            leave_request.save(update_fields=['status', 'reviewed_by', 'reviewed_at', 'rejection_reason', 'updated_by', 'updated_at'])
            ApprovalDocumentService.create_for_leave_request(leave_request, ApprovalDocument.DocumentType.REJECTION, reviewed_by)

            # Notify employee of rejection
            if leave_request.employee.email:
                subject = f"[Leave Rejected] Your {leave_request.leave_type.name} request was rejected"
                body = (
                    f"Hello {leave_request.employee.get_full_name() or leave_request.employee.username},\n\n"
                    f"Your leave request for {leave_request.leave_type.name} from {leave_request.start_date} to {leave_request.end_date} was REJECTED by {reviewed_by.get_full_name() or reviewed_by.username}.\n"
                    f"Reason: {reason.strip() or 'No reason provided.'}\n"
                )
                DeliveryService.enqueue_email(
                    company=leave_request.company,
                    recipient=leave_request.employee.email,
                    subject=subject,
                    body=body,
                    template_name='leave_rejected_notification',
                    sent_by=reviewed_by,
                )
            return leave_request

    @staticmethod
    def request_amendment(leave_request: LeaveRequest, reviewed_by: Employee, reason: str):
        """Return a pending request to its employee for correction without losing audit history."""
        if not reason or not reason.strip():
            raise ValidationError({'reason': 'An amendment reason is required.'})
        with transaction.atomic():
            leave_request = LeaveRequest.objects.select_for_update().get(pk=leave_request.pk)
            if leave_request.status not in (
                LeaveRequest.Status.PENDING_DEPARTMENT_HEAD,
                LeaveRequest.Status.PENDING_HR,
                LeaveRequest.Status.PENDING,
            ):
                raise ValidationError('Only a pending leave request can be amended.')
            step = ApprovalRoutingService.get_current_steps(leave_request).select_for_update().filter(approver=reviewed_by).first()
            if not step:
                raise ValidationError('You are not assigned to the current approval stage.')
            now = timezone.now()
            step.status = LeaveApprovalStep.Status.AMENDMENT_REQUESTED
            step.decision_reason = reason.strip()
            step.decided_at = now
            step.updated_by = reviewed_by
            step.save(update_fields=['status', 'decision_reason', 'decided_at', 'updated_by', 'updated_at'])
            ApprovalDecision.objects.create(
                company=leave_request.company, leave_request=leave_request, approval_step=step,
                actor=reviewed_by, stage=step.stage, sequence=step.sequence,
                decision=ApprovalDecision.Decision.AMENDMENT_REQUESTED, reason=reason.strip(), decided_at=now,
                created_by=reviewed_by, updated_by=reviewed_by,
            )
            leave_request.status = LeaveRequest.Status.AMENDMENT_REQUESTED
            leave_request.amendment_reason = reason.strip()
            leave_request.amendment_requested_by = reviewed_by
            leave_request.amendment_requested_at = now
            leave_request.updated_by = reviewed_by
            leave_request.save(update_fields=[
                'status', 'amendment_reason', 'amendment_requested_by', 'amendment_requested_at', 'updated_by', 'updated_at',
            ])
            if leave_request.employee.email:
                from core.delivery import DeliveryService
                DeliveryService.enqueue_email(
                    company=leave_request.company, recipient=leave_request.employee.email,
                    subject=f'[Leave Amendment Required] {leave_request.leave_type.name}',
                    body=f'Your leave request needs amendment. Reason: {reason.strip()}',
                    template_name='leave_amendment_required', sent_by=reviewed_by,
                )
        return leave_request

    @staticmethod
    def amend_leave(leave_request: LeaveRequest, employee: Employee, **data):
        """Apply employee corrections and start a new, auditable approval round."""
        with transaction.atomic():
            leave_request = LeaveRequest.objects.select_for_update().get(pk=leave_request.pk)
            if leave_request.employee_id != employee.id:
                raise ValidationError('Only the requester may amend this leave request.')
            if leave_request.status != LeaveRequest.Status.AMENDMENT_REQUESTED:
                raise ValidationError('This leave request has not been returned for amendment.')
            start_date, end_date = data['start_date'], data['end_date']
            days = LeaveService.calculate_working_days(start_date, end_date, employee.company)
            if days <= 0:
                raise ValidationError({'start_date': 'Leave must include at least one working day.'})
            if days > settings.MAX_LEAVE_WORKING_DAYS:
                raise ValidationError({'end_date': f'Leave requests may not exceed {settings.MAX_LEAVE_WORKING_DAYS} working days.'})
            document_file = data.pop('document', None)
            if document_file:
                supporting_name = LeaveService.validate_document(document_file)
                document_file.seek(0)
            elif leave_request.leave_type.requires_supporting_document and not leave_request.supporting_zoho_file_id:
                raise ValidationError({'document': f"A supporting document is required for '{leave_request.leave_type.name}' leave requests."})
            else:
                supporting_name = None
            overlap_exists = LeaveRequest.objects.filter(employee=employee, status=LeaveRequest.Status.APPROVED).exclude(pk=leave_request.pk).filter(
                start_date__lte=end_date, end_date__gte=start_date,
            ).exists()
            if overlap_exists:
                raise ValidationError('You already have approved leave during these dates.')
            for field in ('start_date', 'end_date', 'reason', 'contact_during_leave', 'emergency_contact_name', 'emergency_contact_phone', 'handover_contact', 'handover_notes'):
                setattr(leave_request, field, data[field])
            leave_request.days_requested = days
            leave_request.status = LeaveRequest.Status.PENDING_DEPARTMENT_HEAD
            leave_request.approval_round += 1
            leave_request.updated_by = employee
            if document_file:
                folder = OnboardingService.onboard_employee(employee, created_by=employee)
                uploaded = ZohoWorkDriveService().upload_document(
                    employee=employee, folder=folder, document_name=supporting_name,
                    document_type=Path(supporting_name).suffix.lstrip('.'), file_content=document_file.read(), uploaded_by=employee,
                )
                leave_request.supporting_document_name = supporting_name
                leave_request.supporting_zoho_file_id = uploaded.zoho_file_id
                leave_request.supporting_workdrive_url = f'https://workdrive.zoho.com/file/{uploaded.zoho_file_id}'
            leave_request.save()
            ApprovalRoutingService.create_steps(leave_request, approval_round=leave_request.approval_round)
            LeaveService._notify_steps(leave_request, ApprovalRoutingService.get_current_steps(leave_request))
        return leave_request

    @staticmethod
    def cancel_leave(leave_request: LeaveRequest, cancelled_by: Employee):
        if cancelled_by.company_id != leave_request.company_id:
            raise ValidationError('Canceller must belong to the same company.')

        from django.db import transaction

        with transaction.atomic():
            # Lock the leave request row to avoid races with concurrent approvals
            lr = LeaveRequest.objects.select_for_update().get(pk=leave_request.pk)
            if lr.status == LeaveRequest.Status.CANCELLED:
                raise ValidationError('Leave request is already cancelled.')

            from datetime import date
            from core.models import EmployeeRole
            if lr.start_date <= date.today() and cancelled_by.role != EmployeeRole.HR_ADMIN and not cancelled_by.is_superuser:
                raise ValidationError('Leave cannot be cancelled once it has started unless overridden by an HR Admin.')

            was_approved = lr.status == LeaveRequest.Status.APPROVED

            # Mark as cancelled
            lr.status = LeaveRequest.Status.CANCELLED
            lr.updated_by = cancelled_by
            lr.save(update_fields=['status', 'updated_by', 'updated_at'])

            # Record an immutable cancellation decision
            if ApprovalDecision.objects.filter(leave_request=lr, actor=cancelled_by, decision=ApprovalDecision.Decision.CANCELLATION).exists():
                # Already cancelled by this actor
                raise ValidationError('This leave request has already been cancelled by you.')
            ApprovalDecision.objects.create(
                company=lr.company,
                leave_request=lr,
                approval_step=None,
                actor=cancelled_by,
                stage=None,
                sequence=None,
                decision=ApprovalDecision.Decision.CANCELLATION,
                reason='Cancelled via API' if cancelled_by == lr.employee else 'Cancelled by admin',
                decided_at=timezone.now(),
                created_by=cancelled_by,
                updated_by=cancelled_by,
            )

            # If it was approved previously, restore the used_days on the leave balance
            if was_approved:
                balance = LeaveBalance.objects.select_for_update().filter(
                    company=lr.company,
                    employee=lr.employee,
                    leave_type=lr.leave_type,
                ).first()

                if not balance:
                    # If balance row is missing, create one (defensive) and lock it again
                    balance = LeaveBalance.objects.create(
                        company=lr.company,
                        employee=lr.employee,
                        leave_type=lr.leave_type,
                        allocated_days=lr.leave_type.default_days,
                        used_days=0,
                        created_by=cancelled_by,
                        updated_by=cancelled_by,
                    )
                    balance = LeaveBalance.objects.select_for_update().get(pk=balance.pk)

                # Restore used_days
                balance.used_days -= lr.days_requested
                if balance.used_days < 0:
                    # This indicates inconsistent state; abort rather than allow negative balances
                    raise ValidationError('Inconsistent balance state: cannot restore used days because it would be negative.')
                balance.updated_by = cancelled_by
                balance.save(update_fields=['used_days', 'updated_by', 'updated_at'])

            ApprovalDocumentService.create_for_leave_request(lr, ApprovalDocument.DocumentType.CANCELLATION, cancelled_by)

        return lr
